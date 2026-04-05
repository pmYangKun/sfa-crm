# -*- coding: utf-8 -*-
"""从 SFA-CRM-PRD.md 读取内容，生成格式化的 Word 文档。
支持 Mermaid 代码块自动渲染为 PNG 并嵌入文档。"""

import re, json, base64, urllib.request, io, time
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(__file__).parent / "SFA-CRM-PRD.md"
DST = Path(__file__).parent / "SFA-CRM-PRD.docx"
CACHE_DIR = Path(__file__).parent / ".mermaid-cache"

# ------------------------------------------------------------------
# Mermaid 渲染
# ------------------------------------------------------------------

RENDER_SCALE = 4  # 4倍渲染，保证文档放大后依然清晰

def render_mermaid(code: str, index: int) -> bytes | None:
    """使用 mermaid-cli (mmdc) 本地渲染 Mermaid 为高清 PNG。带缓存。"""
    import hashlib, subprocess, tempfile
    CACHE_DIR.mkdir(exist_ok=True)
    cache_key = hashlib.md5(f"{code}__scale{RENDER_SCALE}".encode()).hexdigest()
    cache_file = CACHE_DIR / f"{cache_key}.png"

    if cache_file.exists():
        print(f"  图表 {index}: 使用缓存")
        return cache_file.read_bytes()

    # 写入临时 .mmd 文件
    with tempfile.NamedTemporaryFile(suffix=".mmd", delete=False, mode="w", encoding="utf-8") as f:
        f.write(code)
        mmd_path = f.name

    out_path = mmd_path.replace(".mmd", ".png")

    try:
        result = subprocess.run(
            ["C:/Users/YK/AppData/Roaming/npm/mmdc.cmd", "-i", mmd_path, "-o", out_path, "-s", str(RENDER_SCALE), "-b", "transparent"],
            capture_output=True, text=True, timeout=30
        )
        if Path(out_path).exists() and Path(out_path).stat().st_size > 100:
            data = Path(out_path).read_bytes()
            # 读取实际像素尺寸
            import struct
            w, h = struct.unpack(">II", data[16:24])
            cache_file.write_bytes(data)
            print(f"  图表 {index}: {w}x{h}px, {len(data)//1024}KB")
            return data
        else:
            print(f"  图表 {index}: mmdc 输出为空 - {result.stderr[:200]}")
    except Exception as e:
        print(f"  图表 {index}: 渲染失败 - {e}")
    finally:
        Path(mmd_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    # 降级：尝试 mermaid.ink API
    print(f"  图表 {index}: 降级为 mermaid.ink API...")
    try:
        spec = json.dumps({"code": code, "mermaid": {"theme": "default"}})
        encoded = base64.urlsafe_b64encode(spec.encode()).decode()
        url = f"https://mermaid.ink/img/{encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        data = urllib.request.urlopen(req, timeout=20).read()
        if len(data) > 100:
            cache_file.write_bytes(data)
            print(f"  图表 {index}: API降级成功 ({len(data)//1024}KB)")
            return data
    except Exception as e:
        print(f"  图表 {index}: API降级也失败 - {e}")

    return None


# ------------------------------------------------------------------
# 解析 Markdown
# ------------------------------------------------------------------

def parse_md(path: Path) -> list[dict]:
    """将 Markdown 解析为结构化块列表，正确处理代码块。"""
    text = path.read_text(encoding="utf-8")
    blocks = []
    current_table_lines = []
    in_code_block = False
    code_lang = ""
    code_lines = []

    for line in text.split("\n"):
        # 代码块处理
        stripped = line.strip()

        if stripped.startswith("```") and not in_code_block:
            # 先收集之前的表格
            if current_table_lines:
                blocks.append({"type": "table", "lines": current_table_lines})
                current_table_lines = []
            # 开始代码块
            in_code_block = True
            code_lang = stripped[3:].strip().lower()
            code_lines = []
            continue
        elif stripped.startswith("```") and in_code_block:
            # 结束代码块
            in_code_block = False
            content = "\n".join(code_lines)
            if code_lang == "mermaid":
                blocks.append({"type": "mermaid", "code": content})
            elif content.strip():
                blocks.append({"type": "code", "lang": code_lang, "code": content})
            code_lang = ""
            code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # 表格行
        if stripped.startswith("|"):
            current_table_lines.append(stripped)
            continue
        else:
            if current_table_lines:
                blocks.append({"type": "table", "lines": current_table_lines})
                current_table_lines = []

        if not stripped:
            continue
        elif stripped.startswith("# ") and not stripped.startswith("## "):
            blocks.append({"type": "title", "text": stripped[2:]})
        elif stripped.startswith("## "):
            blocks.append({"type": "h1", "text": stripped[3:]})
        elif stripped.startswith("### "):
            blocks.append({"type": "h2", "text": stripped[4:]})
        elif stripped.startswith("#### "):
            blocks.append({"type": "h3", "text": stripped[5:]})
        elif stripped.startswith("##### "):
            blocks.append({"type": "h4", "text": stripped[6:]})
        elif stripped.startswith("> "):
            blocks.append({"type": "quote", "text": stripped[2:]})
        elif stripped.startswith("---"):
            blocks.append({"type": "hr"})
        elif stripped.startswith("````"):
            continue  # 嵌套代码块标记，跳过
        else:
            blocks.append({"type": "para", "text": stripped})

    if current_table_lines:
        blocks.append({"type": "table", "lines": current_table_lines})

    return blocks


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """解析 Markdown 表格为 (headers, rows)。"""
    if len(lines) < 2:
        return [], []
    headers = [c.strip() for c in lines[0].split("|")[1:-1]]
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    return headers, rows


# ------------------------------------------------------------------
# Word 生成
# ------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_run_font(run, font_name="Arial", east_asia="微软雅黑"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_table(doc, headers, rows):
    if not headers:
        return
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        set_run_font(run)
        set_cell_shading(cell, "D5E8F0")

    for r, row_data in enumerate(rows):
        for c in range(ncols):
            cell = table.rows[r + 1].cells[c]
            text = row_data[c] if c < len(row_data) else ""
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9)
            set_run_font(run)

    doc.add_paragraph()


def build_docx(blocks: list[dict], dst: Path):
    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Heading 样式
    for level, size, color in [(0, 18, "1F4E79"), (1, 14, "2E75B6"), (2, 12, "333333"), (3, 11, "444444")]:
        hs = doc.styles[f"Heading {level + 1}"]
        hs.font.name = "Arial"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor.from_string(color)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 预渲染所有 Mermaid 图表
    mermaid_images = {}
    mermaid_blocks = [(i, b) for i, b in enumerate(blocks) if b["type"] == "mermaid"]
    if mermaid_blocks:
        print(f"\n渲染 {len(mermaid_blocks)} 张 Mermaid 图表...")
        for idx, (block_idx, block) in enumerate(mermaid_blocks):
            img_data = render_mermaid(block["code"], idx + 1)
            if img_data:
                mermaid_images[block_idx] = img_data

    print(f"\n生成 Word 文档...")

    for block_idx, block in enumerate(blocks):
        t = block["type"]

        if t == "title":
            p = doc.add_heading(block["text"], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif t == "h1":
            doc.add_heading(block["text"], level=1)

        elif t == "h2":
            doc.add_heading(block["text"], level=2)

        elif t == "h3":
            doc.add_heading(block["text"], level=3)

        elif t == "h4":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(10.5)
            set_run_font(run)

        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            text = block["text"]
            run = p.add_run(text)
            if "方法论" in text or "💡" in text:
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.size = Pt(9)
            else:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            set_run_font(run)

        elif t == "table":
            headers, rows = parse_table(block["lines"])
            add_table(doc, headers, rows)

        elif t == "mermaid":
            img_data = mermaid_images.get(block_idx)
            if img_data:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(img_data), width=Cm(16))
                doc.add_paragraph()  # spacer
            else:
                # 渲染失败，插入代码块作为后备
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(f"[Mermaid 图表渲染失败，原始代码：]\n{block['code']}")
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        elif t == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["code"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Cm(1)

        elif t == "para":
            text = block["text"]
            p = doc.add_paragraph()
            parts = re.split(r"\*\*(.+?)\*\*", text)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_run_font(run)
                run.font.size = Pt(10.5)
                if i % 2 == 1:
                    run.bold = True

        elif t == "hr":
            doc.add_page_break()

    # 页眉
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("SFA CRM 系统 PRD  |  v1.0  |  2026-04-03")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(dst))
    print(f"\nWord 文档已生成: {dst}")
    print(f"文件大小: {dst.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    blocks = parse_md(SRC)
    mermaid_count = sum(1 for b in blocks if b["type"] == "mermaid")
    print(f"解析到 {len(blocks)} 个内容块，其中 {mermaid_count} 张 Mermaid 图表")
    build_docx(blocks, DST)
